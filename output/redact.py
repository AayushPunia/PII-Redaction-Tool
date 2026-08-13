#!/usr/bin/env python3
"""
PII Redaction Tool for SEBI Red Herring Prospectus
===================================================

Hybrid regex + NER pipeline using Microsoft Presidio with custom recognizers
for Indian corporate identifiers (CIN, PAN, DIN, GSTIN).

Architecture
------------
1. Presidio AnalyzerEngine with spaCy NER (en_core_web_sm) + 7 custom
   recognizers registered for India-specific PII types.
2. Consistent entity mapping (dict: original → fake) using Faker('en_IN')
   so the same PII value maps to the same fake value everywhere.
3. Three-pass .docx processing:
   (a) Normal run-based pass on all paragraphs (top-level + tables +
       headers/footers), with cross-run span handling.
   (b) Field-code pass for hidden HYPERLINK mailto: targets in <w:instrText>.
   (c) Relationship-parts pass for external hyperlink relationships in .rels.

Usage
-----
    python redact.py [--input PATH] [--output PATH] [--log-level DEBUG]

Defaults:
    --input   ./input/Red_Herring_Prospectus.docx
    --output  ./output/redacted_output.docx
"""

import argparse
import logging
import os
import random
import re
import sys
from collections import defaultdict
from typing import Dict, List, Optional, Set, Tuple

from docx import Document
from docx.text.run import Run
from faker import Faker
from lxml import etree
import io
from PIL import Image, ImageDraw, ImageFont

from presidio_analyzer import AnalyzerEngine, EntityRecognizer, RecognizerResult
from presidio_analyzer.nlp_engine import NlpEngineProvider

from recognizers import (
    CINRecognizer,
    CreditCardLuhnRecognizer,
    DINRecognizer,
    DOBRecognizer,
    GSTINRecognizer,
    IndiaPhoneRecognizer,
    PANRecognizer,
)


# ─── Constants ────────────────────────────────────────────────────────────────

LOG = logging.getLogger("pii_redact")

# Word XML namespace
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Hyperlink relationship type URI
HYPERLINK_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


# All entity types the pipeline will detect
SUPPORTED_ENTITIES: List[str] = [
    # NER-based (spaCy)
    "PERSON",
    "LOCATION",
    "ORGANIZATION",
    # Built-in Presidio regex recognizers
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "IP_ADDRESS",
    "US_SSN",
    "CREDIT_CARD",
    # Custom regex recognizers (India-specific)
    "IN_CIN",
    "IN_PAN",
    "IN_DIN",
    "IN_GSTIN",
    "DATE_OF_BIRTH",
]

# Minimum score to accept a detection.  Context-boosted recognizers (DIN, PAN)
# start below this and only cross it when surrounding words confirm the type.
SCORE_THRESHOLD = 0.35

# Regex helpers for field-code pass
FIELD_MAILTO_RE = re.compile(
    r'HYPERLINK\s+"mailto:([^"]+)"', re.IGNORECASE
)
FIELD_HREF_EMAIL_RE = re.compile(
    r'HYPERLINK\s+"([^"]*@[^"]+)"', re.IGNORECASE
)
EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")


# ═══════════════════════════════════════════════════════════════════════════════
# SpacyOrgRecognizer — bridges spaCy's ORG label into Presidio's entity system
# ═══════════════════════════════════════════════════════════════════════════════

class SpacyOrgRecognizer(EntityRecognizer):
    """Maps spaCy ORG NER labels to Presidio's ORGANIZATION entity type.

    Presidio's built-in SpacyRecognizer maps PERSON, GPE/LOC → LOCATION, and
    NRP, but does NOT map ORG by default.  This recognizer fills that gap so
    that company/organization names are included in the analysis results.
    """

    def __init__(self):
        super().__init__(
            supported_entities=["ORGANIZATION"],
            supported_language="en",
            name="SpacyOrgRecognizer",
        )

    def load(self) -> None:
        """No external resources to load."""

    def analyze(
        self,
        text: str,
        entities: List[str],
        nlp_artifacts=None,
        regex_flags=None,
    ) -> List[RecognizerResult]:
        results: List[RecognizerResult] = []
        if nlp_artifacts is None:
            return results
        for ent in nlp_artifacts.entities:
            if ent.label_ == "ORG":
                results.append(
                    RecognizerResult(
                        entity_type="ORGANIZATION",
                        start=ent.start_char,
                        end=ent.end_char,
                        score=0.6,
                    )
                )
        return results


# ═══════════════════════════════════════════════════════════════════════════════
# EntityMapper — consistent PII → fake-value mapping with name-variant coalescing
# ═══════════════════════════════════════════════════════════════════════════════

class EntityMapper:
    """Maintains a deterministic mapping from each unique PII value to a
    Faker-generated replacement.  Same original text → same fake text
    everywhere in the document.

    For PERSON entities, implements local-proximity anaphora resolution:
    - Multi-word names (full names) are registered globally on first sight.
    - Single-word name mentions are resolved to a previously registered full
      name ONLY if that full name also appears in the same local context
      (same paragraph / cell / row).  This avoids the "global substring
      merge" bug where two different people who share a first or last name
      silently collapse into one fake identity.
    - Honorifics (Mr./Mrs./Ms./Dr./Shri/Smt.) are stripped before lookup
      and preserved in the output.
    """

    def __init__(self, seed: int = 42):
        self.faker = Faker("en_IN")
        self.faker.seed_instance(seed)
        random.seed(seed)

        # Main mapping: (entity_type, normalized_lower_text) → fake value
        self._map: Dict[Tuple[str, str], str] = {}

        # Full-name registry for PERSON entities only
        self._full_names: Dict[str, str] = {}  # bare_lower → fake_full_name

        # Per-type redaction counts
        self.stats: Dict[str, int] = defaultdict(int)

    # ── public API ────────────────────────────────────────────────────────

    def get_fake(
        self,
        original: str,
        entity_type: str,
        local_person_names: Optional[Set[str]] = None,
    ) -> str:
        """Return a fake replacement for *original*, creating one if needed."""
        normalized = " ".join(original.split()).strip()
        if not normalized:
            return original

        self.stats[entity_type] += 1

        if entity_type == "PERSON":
            return self._handle_person(normalized, local_person_names)

        key = (entity_type, normalized.lower())
        if key in self._map:
            return self._map[key]

        fake = self._generate(normalized, entity_type)
        self._map[key] = fake
        return fake

    # ── PERSON handling with honorific + variant logic ────────────────────

    @staticmethod
    def _strip_honorific(name: str) -> Tuple[str, str]:
        for prefix in (
            "Mr. ", "Mrs. ", "Ms. ", "Dr. ", "Shri ", "Smt. ",
            "Mr ", "Mrs ", "Ms ", "Dr ",
        ):
            if name.startswith(prefix):
                return prefix, name[len(prefix):]
        return "", name

    def _handle_person(
        self,
        name: str,
        local_names: Optional[Set[str]] = None,
    ) -> str:
        honorific, bare = self._strip_honorific(name)
        bare_lower = bare.lower()

        key = ("PERSON", bare_lower)
        if key in self._map:
            return honorific + self._map[key]

        # ── Local anaphora resolution for single-word mentions ──
        parts = bare.split()
        if len(parts) == 1 and local_names:
            for full_name in local_names:
                _, full_bare = self._strip_honorific(full_name)
                full_parts = full_bare.split()
                if len(full_parts) <= 1:
                    continue
                # Only resolve if our single word matches a part of this
                # full name AND the full name was already mapped.
                for idx, fp in enumerate(full_parts):
                    if fp.lower() == bare_lower:
                        full_key = ("PERSON", full_bare.lower())
                        if full_key in self._map:
                            fake_parts = self._map[full_key].split()
                            if idx < len(fake_parts):
                                resolved = fake_parts[idx]
                            else:
                                resolved = fake_parts[-1]
                            self._map[key] = resolved
                            return honorific + resolved
                        break  # matched but full name not mapped yet

        # ── Generate new fake name ──
        fake_name = self.faker.name()
        # Remove any Faker-generated honorifics to keep output clean
        for prefix in ("Mr. ", "Mrs. ", "Ms. ", "Dr. ", "Shri ", "Smt. "):
            if fake_name.startswith(prefix):
                fake_name = fake_name[len(prefix):]
                break

        self._map[key] = fake_name
        if len(parts) > 1:
            self._full_names[bare_lower] = fake_name

        return honorific + fake_name

    # ── Type-specific fake generators ─────────────────────────────────────

    def _generate(self, original: str, entity_type: str) -> str:
        gen = {
            "EMAIL_ADDRESS": lambda: self.faker.email(),
            "PHONE_NUMBER": lambda: self._fake_phone(original),
            "LOCATION": lambda: f"{self.faker.city()}, {self.faker.state()}",
            "ORGANIZATION": lambda: self.faker.company(),
            "IN_CIN": self._fake_cin,
            "IN_PAN": self._fake_pan,
            "IN_DIN": self._fake_din,
            "IN_GSTIN": self._fake_gstin,
            "CREDIT_CARD": lambda: self.faker.credit_card_number(),
            "US_SSN": lambda: self.faker.ssn(),
            "IP_ADDRESS": lambda: self.faker.ipv4(),
            "DATE_OF_BIRTH": lambda: self.faker.date_of_birth(
                minimum_age=25, maximum_age=75
            ).strftime("%B %d, %Y"),
        }
        return gen.get(entity_type, lambda: "█" * min(len(original), 12))()

    def _fake_phone(self, original: str) -> str:
        """Generate a fake Indian phone number in a format similar to original."""
        if original.lstrip().startswith("+"):
            area = str(random.randint(20, 99))
            num = "".join(str(random.randint(0, 9)) for _ in range(8))
            return f"+91 {area} {num[:4]}{num[4:]}"
        if original.lstrip().startswith("0"):
            area = "0" + str(random.randint(20, 99))
            num = "".join(str(random.randint(0, 9)) for _ in range(8))
            return f"{area}-{num}"
        first = random.choice([6, 7, 8, 9])
        rest = "".join(str(random.randint(0, 9)) for _ in range(9))
        return f"{first}{rest}"

    def _fake_cin(self) -> str:
        p = random.choice("UL")
        d5 = "".join(str(random.randint(0, 9)) for _ in range(5))
        st = random.choice(["MH", "DL", "KA", "TN", "GJ", "WB", "RJ", "UP"])
        yr = str(random.randint(1970, 2024))
        tp = random.choice(["PLC", "PTC"])
        d6 = "".join(str(random.randint(0, 9)) for _ in range(6))
        return f"{p}{d5}{st}{yr}{tp}{d6}"

    def _fake_pan(self) -> str:
        l5 = "".join(random.choice("ABCDEFGHIJKLMNOPQRSTUVWXYZ") for _ in range(5))
        d4 = "".join(str(random.randint(0, 9)) for _ in range(4))
        l1 = random.choice("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
        return f"{l5}{d4}{l1}"

    def _fake_din(self) -> str:
        return "".join(str(random.randint(0, 9)) for _ in range(8))

    def _fake_gstin(self) -> str:
        sc = str(random.randint(1, 37)).zfill(2)
        pan = self._fake_pan()
        en = str(random.randint(1, 9))
        ch = random.choice("ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789")
        return f"{sc}{pan}{en}Z{ch}"


# ═══════════════════════════════════════════════════════════════════════════════
# Run-level helpers
# ═══════════════════════════════════════════════════════════════════════════════

def get_all_runs_ordered(paragraph) -> List[Run]:
    """Return ALL runs in document order, including runs nested inside
    <w:hyperlink> elements (which python-docx's `paragraph.runs` skips).

    This is important because Word can wrap email-address display text in
    a <w:hyperlink> element whose child <w:r> runs are invisible to the
    normal Paragraph.runs property.
    """
    runs: List[Run] = []
    p_elem = paragraph._p  # CT_P element

    for child in p_elem:
        tag = child.tag
        if tag == f"{{{W_NS}}}r":
            runs.append(Run(child, paragraph))
        elif tag == f"{{{W_NS}}}hyperlink":
            for r_elem in child.findall(f"{{{W_NS}}}r"):
                runs.append(Run(r_elem, paragraph))

    return runs


def replace_in_runs(
    runs: List[Run],
    match_start: int,
    match_end: int,
    replacement: str,
) -> None:
    """Replace the character span [match_start, match_end) across *runs*
    with *replacement*, handling cross-run boundaries.

    The entire replacement text goes into the first affected run (preserving
    its formatting).  Subsequent affected runs have their matched portions
    cleared.  Replacements are expected to be applied right-to-left so that
    earlier offsets remain valid.
    """
    # Build offset boundaries per run
    boundaries: List[Tuple[int, int]] = []
    offset = 0
    for run in runs:
        text = run.text or ""
        boundaries.append((offset, offset + len(text)))
        offset += len(text)

    # Identify affected runs (any run whose span overlaps the match)
    affected: List[int] = []
    for i, (rs, re_) in enumerate(boundaries):
        if re_ > match_start and rs < match_end:
            affected.append(i)

    if not affected:
        return

    first_idx = affected[0]
    first_run = runs[first_idx]
    first_start = boundaries[first_idx][0]
    offset_in_first = match_start - first_start

    if len(affected) == 1:
        # Entirely within one run
        offset_end = match_end - first_start
        first_run.text = (
            (first_run.text or "")[:offset_in_first]
            + replacement
            + (first_run.text or "")[offset_end:]
        )
    else:
        # Spans multiple runs
        first_run.text = (first_run.text or "")[:offset_in_first] + replacement

        # Clear middle runs entirely
        for idx in affected[1:-1]:
            runs[idx].text = ""

        # Last run: trim the matched prefix
        last_idx = affected[-1]
        last_run = runs[last_idx]
        last_start = boundaries[last_idx][0]
        offset_in_last = match_end - last_start
        last_run.text = (last_run.text or "")[offset_in_last:]


# ═══════════════════════════════════════════════════════════════════════════════
# Analysis helpers
# ═══════════════════════════════════════════════════════════════════════════════

def is_subject_company(text: str) -> bool:
    """True if *text* contains the subject company name as a case-insensitive substring."""
    return "ksh international" in " ".join(text.lower().split())


def filter_results(
    results: List[RecognizerResult],
    full_text: str,
) -> List[RecognizerResult]:
    """Remove subject-company detections, tiny spans, and sub-threshold hits."""
    filtered: List[RecognizerResult] = []
    for r in results:
        if r.score < SCORE_THRESHOLD:
            continue
        span = full_text[r.start : r.end]
        # Skip subject company
        if r.entity_type in ("ORGANIZATION", "PERSON") and is_subject_company(span):
            continue
        # Skip very short spans (likely noise)
        if len(span.strip()) < 2:
            continue
        filtered.append(r)
    return filtered


def local_person_names(
    results: List[RecognizerResult],
    full_text: str,
) -> Set[str]:
    """Extract multi-word PERSON names from analysis results for local
    anaphora resolution."""
    names: Set[str] = set()
    for r in results:
        if r.entity_type == "PERSON":
            span = full_text[r.start : r.end].strip()
            if len(span.split()) > 1:
                names.add(span)
    return names


# ═══════════════════════════════════════════════════════════════════════════════
# Analyzer factory
# ═══════════════════════════════════════════════════════════════════════════════

def build_analyzer() -> AnalyzerEngine:
    """Construct and return a fully configured Presidio AnalyzerEngine."""
    # ── Ensure spaCy model is available ──
    model_name = "en_core_web_sm"
    try:
        import spacy
        spacy.load(model_name)
    except OSError:
        LOG.info("Downloading spaCy model '%s'...", model_name)
        import subprocess
        subprocess.check_call(
            [sys.executable, "-m", "spacy", "download", model_name],
            stdout=subprocess.DEVNULL,
        )

    # ── Build NLP engine ──
    provider = NlpEngineProvider(
        nlp_configuration={
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": model_name}],
        }
    )
    nlp_engine = provider.create_engine()

    # ── Create analyzer with default recognizers ──
    analyzer = AnalyzerEngine(
        nlp_engine=nlp_engine,
        supported_languages=["en"],
    )

    # ── Register custom recognizers ──
    custom = [
        SpacyOrgRecognizer(),
        CINRecognizer(),
        PANRecognizer(),
        DINRecognizer(),
        GSTINRecognizer(),
        IndiaPhoneRecognizer(),
        CreditCardLuhnRecognizer(),
        DOBRecognizer(),
    ]
    for rec in custom:
        analyzer.registry.add_recognizer(rec)

    LOG.info(
        "Analyzer ready: %d recognizers (%d custom)",
        len(analyzer.registry.recognizers),
        len(custom),
    )
    return analyzer


# ═══════════════════════════════════════════════════════════════════════════════
# Pass 1 — Run-based paragraph redaction
# ═══════════════════════════════════════════════════════════════════════════════

def redact_paragraph(
    paragraph,
    analyzer: AnalyzerEngine,
    mapper: EntityMapper,
    redaction_log: List[dict],
) -> int:
    """Analyse and redact PII in a single paragraph, preserving run formatting.

    Returns the number of entities redacted.
    """
    runs = get_all_runs_ordered(paragraph)
    if not runs:
        return 0

    full_text = "".join(r.text or "" for r in runs)
    if not full_text.strip():
        return 0

    # ── Analyse ──
    try:
        results = analyzer.analyze(
            text=full_text,
            language="en",
            entities=SUPPORTED_ENTITIES,
            score_threshold=SCORE_THRESHOLD,
        )
    except Exception as exc:
        LOG.warning("Analysis error (skipping paragraph): %s", str(exc)[:120])
        return 0

    results = filter_results(results, full_text)
    if not results:
        return 0

    # ── Resolve local names for anaphora ──
    local_names = local_person_names(results, full_text)

    # ── Apply replacements right-to-left to keep offsets valid ──
    results.sort(key=lambda r: r.start, reverse=True)

    count = 0
    for r in results:
        original = full_text[r.start : r.end]
        fake = mapper.get_fake(original, r.entity_type, local_names)
        replace_in_runs(runs, r.start, r.end, fake)

        redaction_log.append(
            {
                "type": r.entity_type,
                "original": original,
                "replacement": fake,
                "score": round(r.score, 2),
            }
        )
        LOG.debug(
            "  [%s] '%s' → '%s' (%.2f)",
            r.entity_type,
            original[:50],
            fake[:50],
            r.score,
        )
        count += 1

    return count


def pass1_run_based(
    doc: Document,
    analyzer: AnalyzerEngine,
    mapper: EntityMapper,
    redaction_log: List[dict],
) -> int:
    """Pass 1: iterate all paragraphs (top-level, tables, headers/footers)
    and redact PII at the run level."""
    total = 0

    # ── 1a. Top-level paragraphs ──
    n_para = len(doc.paragraphs)
    LOG.info("  Top-level paragraphs: %d", n_para)
    for i, para in enumerate(doc.paragraphs):
        if i % 200 == 0 and i > 0:
            LOG.info("    … paragraph %d / %d", i, n_para)
        total += redact_paragraph(para, analyzer, mapper, redaction_log)

    # ── 1b. Tables (with merged-cell dedup and nested-table recursion) ──
    n_tables = len(doc.tables)
    LOG.info("  Tables: %d", n_tables)

    def process_cell(cell):
        """Recursively process a table cell, including nested tables."""
        count = 0
        for para in cell.paragraphs:
            count += redact_paragraph(para, analyzer, mapper, redaction_log)
        for nested_table in cell.tables:
            seen_nested = set()
            for row in nested_table.rows:
                for c in row.cells:
                    if c._tc in seen_nested:
                        continue
                    seen_nested.add(c._tc)
                    count += process_cell(c)
        return count

    for t_idx, table in enumerate(doc.tables):
        seen = set()
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen:
                    continue
                seen.add(cell._tc)
                total += process_cell(cell)

    # ── 1c. Headers and footers ──
    hf_count = 0
    for section in doc.sections:
        for attr in ("header", "footer", "first_page_header", "first_page_footer",
                      "even_page_header", "even_page_footer"):
            try:
                part = getattr(section, attr, None)
                if part is None:
                    continue
                for para in part.paragraphs:
                    hf_count += redact_paragraph(para, analyzer, mapper, redaction_log)
            except Exception:
                pass
    total += hf_count
    LOG.info("  Headers/footers: %d entities redacted", hf_count)

    return total


# ═══════════════════════════════════════════════════════════════════════════════
# Pass 2 — Field-code hyperlinks (<w:instrText>)
# ═══════════════════════════════════════════════════════════════════════════════

def pass2_field_codes(
    doc: Document,
    mapper: EntityMapper,
    redaction_log: List[dict],
) -> int:
    """Find and redact email addresses hidden in Word field-code instructions.

    Word field-code hyperlinks store link targets in <w:instrText> elements
    (e.g., ``HYPERLINK "mailto:someone@example.com"``).  The visible display
    text may differ from the target — both are separate PII surfaces.  Pass 1
    catches the display text; this pass catches the hidden target.
    """
    count = 0

    # Collect all XML roots to search (body + headers/footers)
    roots = [doc.element.body]
    for section in doc.sections:
        for attr in ("header", "footer", "first_page_header", "first_page_footer",
                      "even_page_header", "even_page_footer"):
            try:
                part = getattr(section, attr, None)
                if part is not None and part._element is not None:
                    roots.append(part._element)
            except Exception:
                pass

    for root in roots:
        for instr in root.iter(f"{{{W_NS}}}instrText"):
            text = instr.text
            if not text:
                continue

            # Try mailto: link first
            m = FIELD_MAILTO_RE.search(text)
            if not m:
                m = FIELD_HREF_EMAIL_RE.search(text)
            if m:
                email = m.group(1)
                fake = mapper.get_fake(email, "EMAIL_ADDRESS")
                new_text = text.replace(email, fake)
                if new_text != text:
                    instr.text = new_text
                    count += 1
                    redaction_log.append(
                        {
                            "type": "EMAIL_ADDRESS",
                            "original": email,
                            "replacement": fake,
                            "score": 1.0,
                            "source": "field-code",
                        }
                    )
                    LOG.info("  Field-code: '%s' → '%s'", email, fake)

    return count


# ═══════════════════════════════════════════════════════════════════════════════
# Pass 3 — Relationship-parts hyperlinks (.rels)
# ═══════════════════════════════════════════════════════════════════════════════

def pass3_rels(
    doc: Document,
    mapper: EntityMapper,
    redaction_log: List[dict],
) -> int:
    """Check and redact external hyperlink relationships in document.part.rels.

    Word can also store hyperlinks as relationship entries in the .rels XML
    file.  This pass checks for external hyperlink relationships whose target
    contains an email address or PII-bearing URL, and redacts them.

    For the KSH International RHP, direct inspection of document.xml.rels
    confirms zero r:id-based external hyperlink relationships exist (only
    font/header/footer/image/style relationships).  This pass verifies that
    finding at runtime rather than assuming it.
    """
    count = 0
    try:
        for rel in doc.part.rels.values():
            # Check: is this an external hyperlink?
            is_external = getattr(rel, "is_external", False) or getattr(
                rel, "_is_external", False
            )
            if rel.reltype == HYPERLINK_REL_TYPE and is_external:
                target = getattr(rel, "target_ref", None) or getattr(
                    rel, "_target", None
                )
                if target:
                    m = EMAIL_RE.search(str(target))
                    if m:
                        orig_email = m.group()
                        fake = mapper.get_fake(orig_email, "EMAIL_ADDRESS")
                        # Modify target in-place
                        new_target = str(target).replace(orig_email, fake)
                        rel._target = new_target
                        count += 1
                        redaction_log.append(
                            {
                                "type": "EMAIL_ADDRESS",
                                "original": orig_email,
                                "replacement": fake,
                                "score": 1.0,
                                "source": "rel-part",
                            }
                        )
                        LOG.info("  Rel-part: '%s' → '%s'", orig_email, fake)
    except Exception as exc:
        LOG.warning("Relationship-parts check error: %s", exc)

    LOG.info(
        "  External hyperlink relationships found/redacted: %d "
        "(expected 0 for this document — verified from .rels)",
        count,
    )
    return count


# ═══════════════════════════════════════════════════════════════════════════════
# Pass 4 — Image Redaction (Embedded PAN / Aadhaar cards)
# ═══════════════════════════════════════════════════════════════════════════════

def pass4_images(doc: Document) -> int:
    """Find specific embedded PAN/Aadhaar images and replace them with a solid placeholder."""
    count = 0
    # The original document uses media/image4.png and media/image5.png for the ID cards.
    for rel in doc.part.rels.values():
        if "image" in rel.reltype and rel.target_ref in ["media/image4.png", "media/image5.png"]:
            part = rel.target_part
            if not part:
                continue
            
            try:
                # Get original size to keep the same dimensions
                orig_image = Image.open(io.BytesIO(part.blob))
                width, height = orig_image.size
                
                # Create dark gray placeholder with some text
                placeholder = Image.new("RGB", (width, height), color="#333333")
                draw = ImageDraw.Draw(placeholder)
                
                # We'll just draw a simple cross and some text to indicate redaction
                draw.line((0, 0, width, height), fill="red", width=5)
                draw.line((0, height, width, 0), fill="red", width=5)
                
                # Save to bytes
                out_io = io.BytesIO()
                placeholder.save(out_io, format="PNG")
                
                # Replace the blob
                part._blob = out_io.getvalue()
                count += 1
                LOG.info("  Image redacted: %s (%dx%d)", rel.target_ref, width, height)
            except Exception as e:
                LOG.warning("Failed to redact image %s: %s", rel.target_ref, e)
                
    return count


# ═══════════════════════════════════════════════════════════════════════════════
# Post-run validation
# ═══════════════════════════════════════════════════════════════════════════════

def validate(input_path: str, output_path: str, mapper: EntityMapper) -> bool:
    """Re-open the output .docx and run automated sanity checks."""
    LOG.info("Running validation checks...")
    ok = True

    try:
        inp = Document(input_path)
        out = Document(output_path)
    except Exception as exc:
        LOG.error("Cannot re-open documents for validation: %s", exc)
        return False

    # ── Structure integrity ──
    ip, op = len(inp.paragraphs), len(out.paragraphs)
    it, ot = len(inp.tables), len(out.tables)
    LOG.info("  Paragraphs: input=%d  output=%d  %s", ip, op, "✓" if ip == op else "✗")
    LOG.info("  Tables:      input=%d  output=%d  %s", it, ot, "✓" if it == ot else "✗")
    if ip != op or it != ot:
        ok = False

    # ── Spot checks: PII should be gone ──
    all_out_text = "\n".join(p.text for p in out.paragraphs)
    # Also include table text
    for table in out.tables:
        for row in table.rows:
            for cell in row.cells:
                all_out_text += "\n" + cell.text

    checks = [
        ("Company Secretary name", "Sarthak Malvadkar", False),
        ("CIN (primary)", "U28129PN1979PLC141032", False),
        ("Email (cs.connect)", "cs.connect@kshinternational.com", False),
        ("Email (sarthak)", "Sarthak.malvadkar@kshinterantional.com", False),
        ("Subject company (should survive)", "KSH International Limited", True),
    ]
    for label, needle, should_exist in checks:
        found = needle in all_out_text
        if found == should_exist:
            LOG.info("  %-45s ✓", label)
        else:
            LOG.warning(
                "  %-45s ✗ (%s)",
                label,
                "still present" if found else "missing",
            )
            ok = False

    # ── Consistency: same original → same fake ──
    reverse_map: Dict[Tuple[str, str], Set[str]] = defaultdict(set)
    for (etype, orig), fake in mapper._map.items():
        reverse_map[(etype, orig)].add(fake)
    inconsistent = {k: v for k, v in reverse_map.items() if len(v) > 1}
    if inconsistent:
        LOG.warning("  Consistency: %d originals mapped to multiple fakes!", len(inconsistent))
        ok = False
    else:
        LOG.info("  Consistency: ✓  (%d unique mappings)", len(mapper._map))

    # ── Field-code redaction check ──
    leaked = []
    for instr in out.element.body.iter(f"{{{W_NS}}}instrText"):
        if instr.text and EMAIL_RE.search(instr.text):
            # Check whether the email is one we know about (i.e. still original)
            found_email = EMAIL_RE.search(instr.text).group()
            # If this email is a fake one we generated, that's fine
            fakes = set(v for (et, _), v in mapper._map.items() if et == "EMAIL_ADDRESS")
            if found_email not in fakes:
                leaked.append(found_email)
    if leaked:
        LOG.warning("  Field-code leak: %s", leaked)
        ok = False
    else:
        LOG.info("  Field-code emails in output:                    ✓  all redacted")

    return ok


# ═══════════════════════════════════════════════════════════════════════════════
# Main orchestrator
# ═══════════════════════════════════════════════════════════════════════════════

def process_document(input_path: str, output_path: str) -> Dict:
    """Load → analyse → redact → save → validate."""

    LOG.info("=" * 72)
    LOG.info("PII Redaction Tool — SEBI Red Herring Prospectus")
    LOG.info("=" * 72)
    LOG.info("Input:  %s", os.path.abspath(input_path))
    LOG.info("Output: %s", os.path.abspath(output_path))

    doc = Document(input_path)
    LOG.info("Loaded: %d paragraphs, %d tables", len(doc.paragraphs), len(doc.tables))

    analyzer = build_analyzer()
    mapper = EntityMapper(seed=42)
    redaction_log: List[dict] = []

    # ── Pass 1 ──
    LOG.info("")
    LOG.info("─── Pass 1: Run-based paragraph redaction ─────────────────────")
    p1 = pass1_run_based(doc, analyzer, mapper, redaction_log)
    LOG.info("Pass 1 complete: %d entities redacted", p1)

    # ── Pass 2 ──
    LOG.info("")
    LOG.info("─── Pass 2: Field-code hyperlink redaction ────────────────────")
    p2 = pass2_field_codes(doc, mapper, redaction_log)
    LOG.info("Pass 2 complete: %d field-code hyperlinks redacted", p2)

    # ── Pass 3 ──
    LOG.info("")
    LOG.info("─── Pass 3: Relationship-parts hyperlink check ────────────────")
    p3 = pass3_rels(doc, mapper, redaction_log)
    LOG.info("Pass 3 complete: %d relationship hyperlinks found", p3)

    # ── Pass 4 ──
    LOG.info("")
    LOG.info("─── Pass 4: Embedded Image Redaction ──────────────────────────")
    p4 = pass4_images(doc)
    LOG.info("Pass 4 complete: %d images replaced with placeholders", p4)

    # ── Save ──
    LOG.info("")
    LOG.info("─── Saving output ─────────────────────────────────────────────")
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    LOG.info("Saved: %s", os.path.abspath(output_path))

    # ── Validate ──
    LOG.info("")
    LOG.info("─── Validation ──────────────────────────────────────────────")
    valid = validate(input_path, output_path, mapper)
    LOG.info("Validation: %s", "PASSED" if valid else "ISSUES FOUND")

    # ── Summary ──
    LOG.info("")
    LOG.info("=" * 72)
    LOG.info("SUMMARY")
    LOG.info("=" * 72)
    LOG.info("Pass 1 (run-based):     %d entities", p1)
    LOG.info("Pass 2 (field-codes):   %d hyperlinks", p2)
    LOG.info("Pass 3 (rel-parts):     %d hyperlinks", p3)
    LOG.info("Pass 4 (images):        %d embedded images", p4)
    LOG.info("")
    LOG.info("Detections by entity type:")
    for etype in sorted(mapper.stats):
        LOG.info("  %-22s %d", etype, mapper.stats[etype])
    LOG.info("")
    LOG.info("Unique PII → fake mappings: %d", len(mapper._map))

    return {
        "pass1": p1,
        "pass2": p2,
        "pass3": p3,
        "stats": dict(mapper.stats),
        "mappings": len(mapper._map),
        "log": redaction_log,
    }


# ═══════════════════════════════════════════════════════════════════════════════
# CLI entry point
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(
        description="PII Redaction Tool for SEBI Red Herring Prospectus",
    )
    ap.add_argument(
        "--input", "-i",
        default="./input/Red_Herring_Prospectus.docx",
        help="Path to the input .docx (default: ./input/Red_Herring_Prospectus.docx)",
    )
    ap.add_argument(
        "--output", "-o",
        default="./output/redacted_output.docx",
        help="Path for the redacted output (default: ./output/redacted_output.docx)",
    )
    ap.add_argument(
        "--log-level", "-l",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
    )
    args = ap.parse_args()

    logging.basicConfig(
        level=getattr(logging, args.log_level),
        format="%(asctime)s [%(levelname)s] %(message)s",
        datefmt="%H:%M:%S",
    )

    # Resolve input path (handle underscore vs. space in filename)
    inp = args.input
    if not os.path.isfile(inp):
        alt = inp.replace("Red_Herring_Prospectus", "Red Herring Prospectus")
        if os.path.isfile(alt):
            inp = alt
        else:
            LOG.error("Input file not found: %s", inp)
            sys.exit(1)

    process_document(inp, args.output)


if __name__ == "__main__":
    main()
