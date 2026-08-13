import json
from docx import Document
from collections import defaultdict

with open("draft_gold_standard.json", "r") as f:
    gold = json.load(f)

orig_doc = Document('../input/Red_Herring_Prospectus.docx')
red_doc = Document('./redacted_output.docx')

def extract_all_paragraphs(doc):
    texts = []
    for p in doc.paragraphs:
        if p.text.strip(): texts.append(p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text.strip(): texts.append(p.text)
    return texts

orig_texts = extract_all_paragraphs(orig_doc)
red_texts = extract_all_paragraphs(red_doc)

results = []
metrics = defaultdict(lambda: {"TP": 0, "FP": 0, "FN": 0, "TN": 0})

for g in gold:
    text = g["text"]
    ctx = g["context"]
    g_type = g["type"]
    
    idx = -1
    for i, t in enumerate(orig_texts):
        # Allow partial matches by looking at words or removing newlines
        clean_t = t.replace('\n', ' ')
        if ctx.strip() in clean_t and text in clean_t:
            idx = i
            break
            
    if idx == -1:
        # try matching just by text if context is hard
        for i, t in enumerate(orig_texts):
            clean_t = t.replace('\n', ' ')
            if text in clean_t:
                idx = i
                break
                
    if idx == -1:
        print(f"ERROR: Could not find '{text}' in original document at all!")
        continue
        
    redacted_para = red_texts[idx].replace('\n', ' ')
    orig_para = orig_texts[idx].replace('\n', ' ')
    
    is_present = text in redacted_para
    
    if g_type == "NOT_PII":
        if is_present:
            metrics[g_type]["TN"] += 1
            status = "TN (Preserved)"
        else:
            metrics[g_type]["FP"] += 1
            status = "FP (Improperly Redacted)"
    else:
        if not is_present:
            metrics[g_type]["TP"] += 1
            status = "TP (Redacted)"
        else:
            metrics[g_type]["FN"] += 1
            status = "FN (Leaked)"
            
    results.append({
        "text": text,
        "type": g_type,
        "status": status,
        "original_context": orig_para[:150] + ("..." if len(orig_para) > 150 else ""),
        "redacted_context": redacted_para[:150] + ("..." if len(redacted_para) > 150 else "")
    })

report = ["# PII Redaction Evaluation Report\n"]
report.append("## Overall Metrics by PII Type\n")
report.append("| Entity Type | Precision | Recall | F1 Score | TP | FP | FN | TN |")
report.append("|---|---|---|---|---|---|---|---|")

def safe_div(n, d): return n / d if d else 0.0
overall_tp, overall_fp, overall_fn, overall_tn = 0, 0, 0, 0

for t, m in metrics.items():
    overall_tp += m["TP"]
    overall_fp += m["FP"]
    overall_fn += m["FN"]
    overall_tn += m["TN"]

global_precision = safe_div(overall_tp, overall_tp + overall_fp)
global_recall = safe_div(overall_tp, overall_tp + overall_fn)
global_f1 = safe_div(2 * global_precision * global_recall, global_precision + global_recall)
global_accuracy = safe_div(overall_tp + overall_tn, overall_tp + overall_tn + overall_fp + overall_fn)

for t in sorted(metrics.keys()):
    if t == "NOT_PII": continue
    m = metrics[t]
    precision = safe_div(m["TP"], m["TP"] + m["FP"])
    recall = safe_div(m["TP"], m["TP"] + m["FN"])
    
    # Handle undefined precision for 0/0 cases
    if m["TP"] == 0 and m["FP"] == 0:
        precision_str = "N/A*"
        f1_str = "N/A*"
    else:
        precision_str = f"{precision:.2f}"
        f1 = safe_div(2 * precision * recall, precision + recall)
        f1_str = f"{f1:.2f}"
        
    report.append(f"| {t} | {precision_str} | {recall:.2f} | {f1_str} | {m['TP']} | {m['FP']} | {m['FN']} | {m['TN']} |")

report.append("\n*\* Note: Entity types with 0 true positives and 0 false positives (like IN_DIN) result in a mathematically undefined (0/0) precision. By convention, this represents the absence of predictions rather than a measured metric.*")

report.append("\n## Global Metrics (25-Span Sample)\n")
report.append(f"- **Accuracy (Entity-Level):** {global_accuracy:.2f}")
report.append(f"- **Precision:** {global_precision:.2f}")
report.append(f"- **Recall:** {global_recall:.2f}")
report.append(f"- **F1 Score:** {global_f1:.2f}\n")

report.append("*Note: Accuracy is reported as Entity-Level Exact/Normalized Match Accuracy over the 25-span gold standard sample. Naive full-document token accuracy is not reported as it is artifically inflated (>99%) due to the vast majority of tokens in legal prose being non-PII.*\n")

report.append("## Negative Controls Performance\n")
report.append("These 5 spans test precision by ensuring the tool does **not** over-redact non-PII:\n")
for r in results:
    if r["type"] == "NOT_PII":
        report.append(f"- `{r['text']}`: **{r['status']}**")
        report.append(f"  - *Original:* {r['original_context']}")
        report.append(f"  - *Redacted:* {r['redacted_context']}\n")

report.append("## Positive Controls Performance\n")
for r in results:
    if r["type"] != "NOT_PII":
        report.append(f"- `{r['text']}` ({r['type']}): **{r['status']}**")

with open("evaluation_report.md", "w") as f:
    f.write("\n".join(report))

print("Evaluation report generated.")
